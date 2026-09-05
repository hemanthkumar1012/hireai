import io
import re
from pathlib import Path
from typing import Optional

from docx import Document
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader


MAX_RESUME_SIZE = 5 * 1024 * 1024  # 5 MB

ALLOWED_EXTENSIONS = {".pdf", ".docx"}

ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _clean_text(text: str) -> str:
    """
    Normalize extracted resume text while preserving useful structure.
    """

    if not text:
        return ""

    # Normalize line endings.
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Replace non-breaking spaces.
    text = text.replace("\xa0", " ")

    # Remove excessive spaces/tabs.
    text = re.sub(r"[ \t]+", " ", text)

    # Remove excessive blank lines.
    text = re.sub(r"\n[ \t]*\n[ \t]*\n+", "\n\n", text)

    # Remove spaces at the beginning/end of each line.
    lines = [line.strip() for line in text.split("\n")]

    # Remove empty lines at the beginning/end.
    while lines and not lines[0]:
        lines.pop(0)

    while lines and not lines[-1]:
        lines.pop()

    return "\n".join(lines).strip()


def _extract_pdf_text(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file.
    """

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail="The uploaded PDF could not be read. Please upload a valid PDF file.",
        ) from exc

    if not reader.pages:
        raise HTTPException(
            status_code=400,
            detail="The uploaded PDF does not contain any pages.",
        )

    extracted_pages = []

    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
            if page_text.strip():
                extracted_pages.append(page_text)
        except Exception:
            # Continue processing other pages if one page cannot be extracted.
            continue

    return "\n\n".join(extracted_pages)


def _extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file, including paragraphs and tables.
    """

    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail="The uploaded DOCX file could not be read. Please upload a valid Word document.",
        ) from exc

    sections = []

    # Extract normal paragraphs.
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            sections.append(text)

    # Extract table contents because resumes often contain
    # skills/contact information inside tables.
    for table in document.tables:
        for row in table.rows:
            cells = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    cells.append(cell_text)

            if cells:
                sections.append(" | ".join(cells))

    return "\n".join(sections)


async def extract_resume_text(file: UploadFile) -> str:
    """
    Validate an uploaded resume and extract its text.

    Supported formats:
        - PDF
        - DOCX

    Maximum file size:
        - 5 MB

    Returns:
        Cleaned resume text.
    """

    if not file.filename:
        raise HTTPException(
            status_code=400,
            detail="Please select a resume file.",
        )

    filename = Path(file.filename).name
    extension = Path(filename).suffix.lower()

    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="Unsupported resume format. Please upload a PDF or DOCX file.",
        )

    # Validate MIME type when the browser provides one.
    # Some browsers/proxies may omit or mislabel it, so the extension
    # remains the primary compatibility check.
    if file.content_type and file.content_type not in ALLOWED_CONTENT_TYPES:
        if extension == ".pdf":
            expected_type = "PDF"
        else:
            expected_type = "DOCX"

        raise HTTPException(
            status_code=400,
            detail=f"The uploaded file does not appear to be a valid {expected_type} file.",
        )

    try:
        file_bytes = await file.read()
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail="The resume could not be uploaded. Please try again.",
        ) from exc

    if not file_bytes:
        raise HTTPException(
            status_code=400,
            detail="The uploaded resume is empty.",
        )

    if len(file_bytes) > MAX_RESUME_SIZE:
        raise HTTPException(
            status_code=413,
            detail="Resume file is too large. Maximum allowed size is 5 MB.",
        )

    if extension == ".pdf":
        extracted_text = _extract_pdf_text(file_bytes)
    else:
        extracted_text = _extract_docx_text(file_bytes)

    cleaned_text = _clean_text(extracted_text)

    if not cleaned_text:
        raise HTTPException(
            status_code=422,
            detail=(
                "No readable text could be extracted from this resume. "
                "If this is a scanned PDF, please upload a text-based PDF or DOCX file."
            ),
        )

    # Protect the downstream AI service and database from extremely
    # large extracted text.
    if len(cleaned_text) > 100_000:
        cleaned_text = cleaned_text[:100_000]

    return cleaned_text
